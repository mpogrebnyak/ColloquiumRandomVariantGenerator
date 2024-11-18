import random
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
from collections import defaultdict
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

def read_students(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        students = [' '.join(line.strip().split()) for line in file.readlines() if line.strip()]
    return students


def adjust_column_width(ws):
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column].width = adjusted_width


def fill_alternate_rows(ws):
    fill = PatternFill(start_color="ADD8E6", end_color="ADD8E6", fill_type="solid")
    for i, row in enumerate(ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column), start=2):
        if i % 2 == 0:
            for cell in row:
                cell.fill = fill


def read_questions_from_excel(file_path):
    df = pd.read_excel(file_path, header=None)
    questions = {index + 1: row[0] for index, row in df.iterrows()}
    return questions


def add_page_break(document):
    page_break = OxmlElement('w:br')
    page_break.set(qn('w:type'), 'page')
    document.add_paragraph()._element.append(page_break)


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        left={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        right={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
    )
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    for border_name in ["top", "left", "bottom", "right"]:
        if border_name in kwargs:
            border = OxmlElement(f"w:{border_name}")
            for key, value in kwargs[border_name].items():
                border.set(qn(f"w:{key}"), str(value))
            tcPr.append(border)


def create_word_document(students_questions, questions_text, output_path):
    document = Document()

    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "Коллоквиум по курсу \"Основы программирования\""
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    for section in document.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.text = "Преподаватель: Погребняк Максим Анатольевич"
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    count = 1
    for student, question_numbers in students_questions.items():
        document.add_paragraph()

        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title_paragraph.add_run("Бланк заданий/ответов №" + str(count))
        run.bold = True
        run.font.size = Pt(16)
        count = count + 1

        fio_group_paragraph = document.add_paragraph()
        fio_group_paragraph.add_run("ФИО: _______________________________________  Группа: ____________").bold = True
        fio_group_paragraph.paragraph_format.space_after = Pt(5)

        table = document.add_table(rows=2, cols=12)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False

        for i in range(10):
            cell = table.cell(0, i)
            cell.text = str(i + 1)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table.cell(0, 10).text = "к/р"
        table.cell(0, 10).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table.cell(0, 11).text = "всего"
        table.cell(0, 11).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for i in range(12):
            table.cell(1, i).text = ""
            table.cell(1, i).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, top={"sz": 6, "val": "single", "color": "000000"},
                                bottom={"sz": 6, "val": "single", "color": "000000"},
                                left={"sz": 6, "val": "single", "color": "000000"},
                                right={"sz": 6, "val": "single", "color": "000000"})

        explanation_text = "Таблица заполняется преподавателем. Самостоятельно заполнять её не нужно. " \
                           "За каждое задание можно получить от 0 до 1 балла.\n" \
                           "Ответы на задания необходимо писать на этом бланке с передней и обратной стороны.\n" \
                           "В поле 'к/р' указываются " \
                           "баллы за контрольные задания с практики, рассчитывающиеся по формуле: решённые задачи / задания * 2.\n" \
                           "Для успешной сдачи коллоквиума необходимо в сумме набрать более 7.5 баллов."

        explanation_paragraph = document.add_paragraph(explanation_text)
        explanation_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = explanation_paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(8)
        explanation_paragraph.paragraph_format.space_after = Pt(0)
        explanation_paragraph.paragraph_format.space_before = Pt(0)

        document.add_paragraph()

        # Add questions with numbering
        for i, q_num in enumerate(question_numbers, start=1):
            paragraph = document.add_paragraph(f"{i}. {questions_text[q_num]}")
            paragraph.paragraph_format.space_after = Pt(5)  # Add space after each question
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(9)

        add_page_break(document)

    document.save(output_path)


def main():
    students_file = 'students.txt'
    questions_file = 'questions.xlsx'
    output_path = 'students_assignments.docx'  # Single output document

    students = read_students(students_file)
    students.sort()

    num_students = len(students)
    total_questions = 250
    questions_per_student = 10

    if num_students * questions_per_student > total_questions * 3:
        raise ValueError("Количество студентов превышает возможное количество вопросов с учетом повторений.")

    questions = list(range(1, total_questions + 1))
    question_count = defaultdict(int)
    students_questions = defaultdict(list)

    for student in students:
        assigned_questions = set()
        while len(assigned_questions) < questions_per_student:
            question = random.choice(questions)
            if question_count[question] < 3 and question not in assigned_questions:
                assigned_questions.add(question)
                question_count[question] += 1
        students_questions[student] = list(assigned_questions)

    df = pd.DataFrame.from_dict(students_questions, orient='index',
                                columns=[f'Question {i + 1}' for i in range(questions_per_student)])

    df.index.name = 'Student'
    df.reset_index(inplace=True)

    excel_path = "students_questions.xlsx"
    df.to_excel(excel_path, index=False)

    wb = load_workbook(excel_path)
    ws = wb.active

    adjust_column_width(ws)
    fill_alternate_rows(ws)
    wb.save(excel_path)

    questions_text = read_questions_from_excel(questions_file)
    create_word_document(students_questions, questions_text, output_path)

    print("Вопросы успешно распределены и сохранены в файл students_questions.xlsx")
    print(f"Документ Word для студентов успешно создан и сохранен в файл {output_path}")


if __name__ == "__main__":
    main()
