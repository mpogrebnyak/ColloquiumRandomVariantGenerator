from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_page_break(document):
    page_break = OxmlElement('w:br')
    page_break.set(qn('w:type'), 'page')
    document.add_paragraph()._element.append(page_break)


def set_cell_border(cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for border_name in ["top", "left", "bottom", "right"]:
        if border_name in kwargs:
            border = OxmlElement(f"w:{border_name}")
            for key, value in kwargs[border_name].items():
                border.set(qn(f"w:{key}"), str(value))
            tcPr.append(border)

def set_document_spacing(document):

    for paragraph in document.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(4)


def create_word_document(students_questions, output_path):
    document = Document()

    # Header and footer setup
    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = "Коллоквиум по курсу \"Основы программирования\""
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.text = "Преподаватель: Погребняк Максим Анатольевич"
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Narrow margins
    for section in document.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    for count, (student, questions) in enumerate(students_questions.items(), start=1):
        document.add_paragraph()

        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title_paragraph.add_run(f"Бланк заданий/ответов №{count}")
        run.bold = True
        run.font.size = Pt(16)

        fio_group_paragraph = document.add_paragraph()
        fio_group_paragraph.add_run("ФИО: _______________________________________  Группа: ____________").bold = True

        table = document.add_table(rows=2, cols=11)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False

        for i in range(10):
            cell = table.cell(0, i)
            cell.text = str(i + 1)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        table.cell(0, 10).text = "Сумма"
        table.cell(0, 10).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, top={"sz": 6, "val": "single", "color": "000000"},
                                bottom={"sz": 6, "val": "single", "color": "000000"},
                                left={"sz": 6, "val": "single", "color": "000000"},
                                right={"sz": 6, "val": "single", "color": "000000"})
        explanation_text = (
            "Таблицу заполняет преподаватель. Самостоятельно заполнять её не нужно. За теоретическое задание можно получить от 0 до 1 балла, "
            "за алгоритмическое — от 0 до 2 баллов (1 балл — за алгоритм и оценку сложности, 1 балл — за реализацию). Ответы необходимо писать на этом бланке с обеих сторон."
        )

        explanation_paragraph = document.add_paragraph()
        explanation_run = explanation_paragraph.add_run(explanation_text)
        explanation_run.font.size = Pt(9)
        explanation_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        theory_paragraph = document.add_paragraph()
        theory_run = theory_paragraph.add_run("Теоретическая секция:")
        theory_run.bold = True
        theory_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        theory_questions = [q for q in questions if q[1] == "Theory"]
        for i, (question, _) in enumerate(theory_questions, start=1):
            document.add_paragraph(f"{i}. {question}").paragraph_format.line_spacing = Pt(10)

        add_page_break(document)

        practical_paragraph = document.add_paragraph()
        practical_run = practical_paragraph.add_run("Алгоритмическая секция:")
        practical_run.bold = True
        practical_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        practical_questions = [q for q in questions if q[1] == "Practical"]
        for i, (question, _) in enumerate(practical_questions, start=1):
            document.add_paragraph(f"{i+len(theory_questions)}. {question}").paragraph_format.line_spacing = Pt(10)
            document.add_paragraph()

        add_page_break(document)

    set_document_spacing(document)
    document.save(output_path)
